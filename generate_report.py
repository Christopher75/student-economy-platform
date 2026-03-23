"""
Script to generate Report.docx for the Student Economy Platform project.
Run: python generate_report.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_heading_color(paragraph, rgb):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*rgb)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2563EB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def create_report():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # =====================================================================
    # TITLE PAGE
    # =====================================================================
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('\n\nSTUDENT ECONOMY PLATFORM')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run('A Campus Marketplace & Skill Exchange Web Application')
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x60, 0x74, 0x8A)

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    meta_info = [
        ('Institution', 'Cavendish University Uganda'),
        ('Programme', 'Information Technology / Computer Science'),
        ('Module', 'Group Project / Software Development'),
        ('Academic Year', '2024/2025'),
        ('Date', datetime.date.today().strftime('%B %d, %Y')),
    ]

    for label, value in meta_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f'{label}: ')
        r1.bold = True
        r1.font.size = Pt(12)
        r2 = p.add_run(value)
        r2.font.size = Pt(12)

    doc.add_page_break()

    # =====================================================================
    # 1. PROJECT OVERVIEW
    # =====================================================================
    h = doc.add_heading('1. Project Overview', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    doc.add_paragraph(
        'The Student Economy Platform (SEP) is a campus-exclusive web application built using '
        'the Django framework that combines two powerful modules into one unified digital ecosystem '
        'for university students. The platform integrates a Campus Marketplace — where verified '
        'university students can buy and sell used items such as textbooks, electronics, furniture, '
        'and clothing — with a Campus Skill Exchange — where students can offer and hire peer services '
        'such as tutoring, graphic design, IT support, photography, writing, and translation.'
    )

    doc.add_paragraph(
        'The platform is built on the entrepreneurship concept of Synthesis: the idea that combining '
        'two existing, proven concepts (a peer-to-peer marketplace and a freelance services marketplace) '
        'into a single student-focused platform creates something more valuable than either idea alone. '
        'Both modules serve the same user (a verified university student), share the same trust '
        'infrastructure (student ID verification and ratings), and pursue the same deeper mission: '
        'putting economic power back into students\' hands — enabling them to earn, save, and transact '
        'within their own campus community.'
    )

    # =====================================================================
    # 2. ENTREPRENEURSHIP CONCEPT
    # =====================================================================
    h = doc.add_heading('2. Entrepreneurship Concept: Synthesis & the Student Economy', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    doc.add_heading('2.1 The Synthesis Principle', 2)
    doc.add_paragraph(
        'Synthesis, as an entrepreneurship concept, refers to the deliberate combination of two or more '
        'existing ideas, markets, or systems to create a new entity that is greater than the sum of its '
        'parts. Historical examples include the smartphone (phone + computer + camera + GPS), Airbnb '
        '(hotels + the internet + spare rooms), and Uber (taxis + smartphones + GPS).'
    )
    doc.add_paragraph(
        'The Student Economy Platform applies Synthesis by merging the peer-to-peer marketplace model '
        '(popularised by platforms like Facebook Marketplace, OLX, and Jiji) with the gig economy model '
        '(popularised by platforms like Fiverr and Upwork), and restricting both to a single, verified '
        'campus community. The result is a trusted, contextual, student-first platform that neither '
        'concept would deliver on its own.'
    )

    doc.add_heading('2.2 The Student Economy Vision', 2)
    doc.add_paragraph(
        'University students in Uganda and across East Africa face a unique economic reality: they have '
        'significant skills and assets (textbooks, electronics, time, knowledge), but lack accessible, '
        'trusted channels to monetise or exchange them efficiently. Off-campus marketplaces are '
        'untrustworthy for student transactions. Generic freelance platforms are not designed for '
        'short-term campus-based arrangements. The Student Economy Platform fills this gap by creating '
        'a micro-economy within the campus gates — where every seller, buyer, provider, and client is '
        'a fellow student with a verifiable identity.'
    )

    # =====================================================================
    # 3. SETUP INSTRUCTIONS
    # =====================================================================
    h = doc.add_heading('3. Setup & Installation Instructions', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    doc.add_heading('3.1 Prerequisites', 2)
    prereqs = ['Python 3.10 or higher', 'pip (Python package manager)', 'Git (optional, for cloning)']
    for p in prereqs:
        doc.add_paragraph(p, style='List Bullet')

    doc.add_heading('3.2 Step-by-Step Installation', 2)

    steps = [
        ('Clone or Download the Project',
         'git clone <repository-url>\ncd "Student Economy Platform"'),
        ('Create a Virtual Environment',
         'python -m venv venv\n\n# Windows:\nvenv\\Scripts\\activate\n\n# macOS/Linux:\nsource venv/bin/activate'),
        ('Install Dependencies',
         'pip install -r requirements.txt'),
        ('Configure Environment Variables',
         'Copy or rename .env.example to .env (or edit the existing .env file).\n'
         'Key settings:\n'
         '  SECRET_KEY=your-secret-key-here\n'
         '  DEBUG=True\n'
         '  ALLOWED_HOSTS=localhost,127.0.0.1'),
        ('Apply Database Migrations',
         'python manage.py migrate'),
        ('Seed Demo Data',
         'python manage.py seed_data\n\nTo reset and reseed:\npython manage.py seed_data --clear'),
        ('Run the Development Server',
         'python manage.py runserver\n\nThe platform is now available at: http://127.0.0.1:8000/'),
    ]

    for i, (title_text, code) in enumerate(steps, 1):
        p = doc.add_paragraph()
        r = p.add_run(f'Step {i}: {title_text}')
        r.bold = True

        code_para = doc.add_paragraph()
        code_run = code_para.add_run(code)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(9)
        code_run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        code_para.paragraph_format.left_indent = Inches(0.4)

    # =====================================================================
    # 4. FEATURES BUILT
    # =====================================================================
    h = doc.add_heading('4. Features Built', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    feature_groups = [
        ('Authentication & User System', [
            'Custom user model with email-based login',
            'Registration with: full name, university email, student ID, university, course, year of study, phone number, profile photo',
            'Email verification on registration (console output in development)',
            'Password reset via email',
            'User profile pages with listings, skills, reviews, and bio',
            'Profile editing',
            'Student verification system (admin-controlled)',
        ]),
        ('Campus Marketplace', [
            'Create, edit, and delete listings (owner-only)',
            'Listings with: title, description, category (8 categories), condition (5 levels), price, negotiable toggle, campus location',
            'Photo upload: up to 5 photos per listing with primary photo selection',
            'Browse all listings with pagination (12 per page)',
            'Filter by category, condition, price range, and university',
            'Search by keyword across title and description',
            'Sort by newest, price (low-high, high-low), and most viewed',
            'Mark item as sold',
            'Save/bookmark listings (wishlist)',
            'View wishlist / saved listings',
            'Contact seller via internal messaging system',
            'Report a listing as inappropriate',
            'Related listings shown on detail page',
            'Listing view counter',
            'Featured listings for homepage',
        ]),
        ('Campus Skill Exchange', [
            'Create, edit, and delete skill offerings (provider-only)',
            'Skill offerings with: title, description, category (10 categories), delivery method (in-person/online/both), price type, price range, estimated duration, availability',
            'Portfolio items per skill (images and/or links)',
            'Browse and search skill offerings',
            'Filter by category, delivery method, price, and university',
            'Book a skill session (booking request with date and notes)',
            'Booking management: pending / confirmed / completed / cancelled',
            'Provider actions: accept, decline, complete bookings',
            'Post-completion reviews with 1-5 star ratings and comments',
            'Ratings aggregated and displayed on skill listings and provider profiles',
        ]),
        ('Messaging System', [
            'Internal messaging between users (not email)',
            'Conversations linked to marketplace listings or skill bookings',
            'Inbox page with all conversations sorted by recent activity',
            'Unread message count in the navbar',
            'Chat-style conversation interface with auto-scroll',
            'Enter to send (Shift+Enter for new line)',
            'Mark messages as read on conversation open',
        ]),
        ('Notifications System', [
            'In-platform notification bell with unread badge in navbar',
            'Notifications for: new messages, booking requests, booking accepted/declined/completed, item marked as sold, new review received',
            'Mark individual notifications as read (AJAX)',
            'Mark all notifications as read',
            'Recent notifications dropdown in navbar',
        ]),
        ('Smart / Creative Features', [
            'Trust Score: computed from email verification, admin verification, and completed transactions',
            'Verified Student badge shown on profiles and listings',
            'Activity Status: "Active today", "Active this week", "Inactive"',
            'Hot Listings: most-viewed listings featured on homepage',
            'Featured Listings: admin-markable listings for homepage spotlight',
            'Listing view counter tracking',
            'User dashboard with personal analytics (listing views, booking stats)',
            'Platform statistics on homepage (total users, listings, skills)',
            'Seller response and profile rating system',
        ]),
        ('Admin Dashboard', [
            'Customised Django admin panel',
            'Admin branding: "Student Economy Platform — Admin"',
            'Full user management with approve/reject student accounts',
            'All marketplace models with filtering and search',
            'Listing report management with resolve action',
            'Skill offering and booking management',
            'Review and messaging oversight',
            'Notification management',
        ]),
        ('User Experience', [
            'Mobile-responsive Bootstrap 5 design',
            'Custom CSS with consistent blue (#2563eb) primary colour scheme',
            'Animated card hover effects',
            'Flash message alerts (auto-dismiss after 5 seconds)',
            'Custom 404 and 500 error pages',
            'Empty state pages with helpful messages and actions',
            'Image upload preview (client-side JavaScript)',
            'Interactive star rating widget for review submission',
            'Photo gallery with thumbnail switching on listing detail pages',
            'Offcanvas filter panel on mobile',
        ]),
    ]

    for group_title, features in feature_groups:
        doc.add_heading(group_title, 2)
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')

    # =====================================================================
    # 5. URL MAP
    # =====================================================================
    h = doc.add_heading('5. Application URLs / Pages', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    url_table = doc.add_table(rows=1, cols=3)
    url_table.style = 'Light List Accent 1'
    hdr_cells = url_table.rows[0].cells
    hdr_cells[0].text = 'URL'
    hdr_cells[1].text = 'Page Name'
    hdr_cells[2].text = 'Access'

    urls = [
        ('/', 'Homepage', 'Public'),
        ('/accounts/register/', 'Student Registration', 'Public'),
        ('/accounts/login/', 'Login', 'Public'),
        ('/accounts/logout/', 'Logout', 'Authenticated'),
        ('/accounts/verify-email/<token>/', 'Email Verification', 'Public'),
        ('/accounts/profile/<username>/', 'User Profile', 'Public'),
        ('/accounts/profile/edit/', 'Edit Profile', 'Authenticated'),
        ('/accounts/dashboard/', 'My Dashboard', 'Authenticated'),
        ('/accounts/password-change/', 'Change Password', 'Authenticated'),
        ('/accounts/password-reset/', 'Password Reset', 'Public'),
        ('/marketplace/', 'Browse Listings', 'Public'),
        ('/marketplace/create/', 'Create Listing', 'Authenticated'),
        ('/marketplace/<pk>/', 'Listing Detail', 'Public'),
        ('/marketplace/<pk>/edit/', 'Edit Listing', 'Owner'),
        ('/marketplace/<pk>/delete/', 'Delete Listing', 'Owner'),
        ('/marketplace/<pk>/mark-sold/', 'Mark as Sold', 'Owner'),
        ('/marketplace/<pk>/save/', 'Toggle Wishlist', 'Authenticated'),
        ('/marketplace/<pk>/report/', 'Report Listing', 'Authenticated'),
        ('/marketplace/wishlist/', 'My Wishlist', 'Authenticated'),
        ('/skills/', 'Browse Skills', 'Public'),
        ('/skills/create/', 'Create Skill Offering', 'Authenticated'),
        ('/skills/<pk>/', 'Skill Detail', 'Public'),
        ('/skills/<pk>/edit/', 'Edit Skill Offering', 'Provider'),
        ('/skills/<pk>/delete/', 'Delete Skill Offering', 'Provider'),
        ('/skills/<pk>/book/', 'Book Skill Session', 'Authenticated'),
        ('/bookings/', 'My Bookings', 'Authenticated'),
        ('/bookings/<pk>/', 'Booking Detail', 'Participants'),
        ('/bookings/<pk>/accept/', 'Accept Booking', 'Provider'),
        ('/bookings/<pk>/decline/', 'Decline Booking', 'Provider'),
        ('/bookings/<pk>/complete/', 'Mark Complete', 'Provider'),
        ('/bookings/<pk>/review/', 'Leave Review', 'Participants'),
        ('/messages/', 'Inbox', 'Authenticated'),
        ('/messages/<pk>/', 'Conversation', 'Participants'),
        ('/messages/start/<username>/', 'Start Conversation', 'Authenticated'),
        ('/messages/start/listing/<pk>/', 'Message About Listing', 'Authenticated'),
        ('/notifications/', 'Notifications', 'Authenticated'),
        ('/notifications/<pk>/read/', 'Mark Notification Read', 'Authenticated'),
        ('/notifications/mark-all-read/', 'Mark All Read', 'Authenticated'),
        ('/admin/', 'Admin Dashboard', 'Staff/Admin'),
    ]

    for url, name, access in urls:
        row = url_table.add_row()
        row.cells[0].text = url
        row.cells[1].text = name
        row.cells[2].text = access

    # =====================================================================
    # 6. DEFAULT CREDENTIALS
    # =====================================================================
    h = doc.add_heading('6. Default Login Credentials (After Seeding)', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    creds_table = doc.add_table(rows=1, cols=4)
    creds_table.style = 'Light List Accent 1'
    hdr = creds_table.rows[0].cells
    hdr[0].text = 'Role'
    hdr[1].text = 'Email'
    hdr[2].text = 'Password'
    hdr[3].text = 'University'

    credentials = [
        ('Admin', 'admin@studenteconomy.ug', 'admin1234', 'N/A'),
        ('Student (CS)', 'amara.nakato@cavendish.ac.ug', 'student1234', 'Cavendish University Uganda'),
        ('Student (Commerce)', 'brian.omoding@mak.ac.ug', 'student1234', 'Makerere University'),
        ('Student (IT)', 'sarah.achieng@must.ac.ug', 'student1234', 'MUST'),
        ('Student (Media)', 'david.ssempala@iuea.ac.ug', 'student1234', 'IUEA'),
        ('Student (Nursing)', 'grace.akello@kcca.ac.ug', 'student1234', 'KCCA University'),
        ('Student (Eng)', 'peter.okello@nkumba.ac.ug', 'student1234', 'Nkumba University'),
        ('Student (BBA)', 'fatima.namutebi@cavendish.ac.ug', 'student1234', 'Cavendish University Uganda'),
        ('Student (Law)', 'kenneth.mwesigwa@mak.ac.ug', 'student1234', 'Makerere University'),
    ]

    for role, email, pw, uni in credentials:
        row = creds_table.add_row()
        row.cells[0].text = role
        row.cells[1].text = email
        row.cells[2].text = pw
        row.cells[3].text = uni

    # =====================================================================
    # 7. TECHNICAL STACK
    # =====================================================================
    h = doc.add_heading('7. Technical Stack', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Light List Accent 1'
    hdr = tech_table.rows[0].cells
    hdr[0].text = 'Component'
    hdr[1].text = 'Technology'
    hdr[2].text = 'Version / Notes'

    tech = [
        ('Backend Framework', 'Django', '6.0.3'),
        ('Database (Dev)', 'SQLite', 'File-based, zero-config'),
        ('Database (Prod-ready)', 'PostgreSQL', 'Change ENGINE in settings'),
        ('Frontend', 'Bootstrap 5', 'CDN-loaded, responsive'),
        ('Icons', 'Bootstrap Icons', 'CDN-loaded'),
        ('Image Processing', 'Pillow', '12.1.1'),
        ('Form Rendering', 'django-crispy-forms + crispy-bootstrap5', '2.6 / 2026.3'),
        ('Template Tweaks', 'django-widget-tweaks', '1.5.1'),
        ('Env Management', 'python-decouple', '3.8'),
        ('File Cleanup', 'django-cleanup', '9.0.0'),
        ('Language', 'Python', '3.12.7'),
        ('Authentication', 'Django Built-in (extended)', 'Email as USERNAME_FIELD'),
    ]

    for comp, tech_name, notes in tech:
        row = tech_table.add_row()
        row.cells[0].text = comp
        row.cells[1].text = tech_name
        row.cells[2].text = notes

    # =====================================================================
    # 8. KNOWN LIMITATIONS & FUTURE IMPROVEMENTS
    # =====================================================================
    h = doc.add_heading('8. Known Limitations & Future Improvements', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    doc.add_heading('Current Limitations', 2)
    limitations = [
        'Messaging is page-refresh based (not real-time). Upgrading to WebSockets with Django Channels would enable real-time chat.',
        'Media files are stored locally. Production deployment should use cloud storage (AWS S3, Cloudinary, or similar).',
        'Email is sent to the console in development mode. A production SMTP server must be configured for real email delivery.',
        'Search uses simple Django ORM ILIKE queries. Full-text search with PostgreSQL full-text or Elasticsearch would improve relevance.',
        'No mobile app — platform is web-only. A React Native or Flutter app would extend reach significantly.',
        'No payment processing integration. UG-relevant options (MTN Mobile Money, Airtel Money) are planned for version 2.',
        'Booking does not include a proper availability calendar — users describe availability in free text.',
        'No image moderation. Admin must manually review flagged listings.',
    ]
    for lim in limitations:
        doc.add_paragraph(lim, style='List Bullet')

    doc.add_heading('Planned Future Improvements', 2)
    future = [
        'Real-time messaging with Django Channels and Redis',
        'Mobile Money payment integration (MTN MoMo, Airtel Money) for in-platform transactions',
        'Push notifications (web push API)',
        'Advanced search with filters, tags, and full-text indexing',
        'Automated student ID verification via university system API integration',
        'Calendar-based availability booking system for skills',
        'Progressive Web App (PWA) capability for mobile-first experience',
        'Multi-language support (Luganda, Swahili, French for wider East African reach)',
        'Escrow system for secure transactions',
        'Alumni network extension for post-graduation mentoring',
    ]
    for f in future:
        doc.add_paragraph(f, style='List Bullet')

    # =====================================================================
    # 9. PROJECT STRUCTURE
    # =====================================================================
    h = doc.add_heading('9. Project Structure', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    structure = doc.add_paragraph()
    structure_run = structure.add_run(
        'Student Economy Platform/\n'
        '├── student_economy/        # Django project configuration\n'
        '│   ├── settings.py         # All settings\n'
        '│   └── urls.py             # Root URL configuration\n'
        '├── accounts/               # User auth & profiles\n'
        '├── marketplace/            # Campus Marketplace module\n'
        '├── skills/                 # Campus Skill Exchange module\n'
        '├── messaging/              # Internal messaging system\n'
        '├── notifications/          # In-platform notifications\n'
        '├── core/                   # Homepage, utilities, error pages\n'
        '│   └── management/commands/seed_data.py\n'
        '├── templates/              # All HTML templates\n'
        '├── static/css/main.css     # Custom CSS\n'
        '├── static/js/main.js       # Custom JavaScript\n'
        '├── media/                  # User-uploaded files\n'
        '├── requirements.txt        # Python dependencies\n'
        '├── .env                    # Environment variables\n'
        '├── PROJECT_PLAN.md         # Architecture plan\n'
        '├── DECISIONS.md            # Design decision log\n'
        '└── db.sqlite3              # SQLite database\n'
    )
    structure_run.font.name = 'Courier New'
    structure_run.font.size = Pt(9)

    # =====================================================================
    # 10. TEAM CREDITS
    # =====================================================================
    h = doc.add_heading('10. Team Credits', 1)
    set_heading_color(h, (0x25, 0x63, 0xEB))

    doc.add_paragraph(
        'This project was developed as a group assignment by IT students at Cavendish University Uganda. '
        'All team members contributed to the design, development, testing, and documentation of the platform.'
    )

    team_table = doc.add_table(rows=1, cols=3)
    team_table.style = 'Light List Accent 1'
    hdr = team_table.rows[0].cells
    hdr[0].text = 'Team Member'
    hdr[1].text = 'Student ID'
    hdr[2].text = 'Primary Contribution'

    team = [
        ('Team Member 1', 'CUU/XXX/2021/001', 'Project Lead — Architecture design, Django project setup, accounts app, authentication system'),
        ('Team Member 2', 'CUU/XXX/2021/002', 'Backend Developer — Marketplace app, listing models, photo upload, search and filter'),
        ('Team Member 3', 'CUU/XXX/2021/003', 'Backend Developer — Skills app, booking system, review and rating system'),
        ('Team Member 4', 'CUU/XXX/2021/004', 'Backend Developer — Messaging system, notifications, context processors'),
        ('Team Member 5', 'CUU/XXX/2021/005', 'Frontend Developer — Base template, homepage, Bootstrap 5 integration, responsive design'),
        ('Team Member 6', 'CUU/XXX/2021/006', 'Frontend Developer — Marketplace templates, skill templates, booking UI, photo gallery'),
        ('Team Member 7', 'CUU/XXX/2021/007', 'Frontend Developer — Messaging UI, notifications UI, user profile and dashboard templates'),
        ('Team Member 8', 'CUU/XXX/2021/008', 'QA & Documentation — Testing, admin customisation, seed data, project documentation'),
    ]

    for name, sid, contrib in team:
        row = team_table.add_row()
        row.cells[0].text = name
        row.cells[1].text = sid
        row.cells[2].text = contrib

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f'Student Economy Platform | Cavendish University Uganda | {datetime.date.today().year}'
    )
    footer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    footer_run.font.size = Pt(9)

    doc.save('Report.docx')
    print('Report.docx generated successfully.')


if __name__ == '__main__':
    create_report()
